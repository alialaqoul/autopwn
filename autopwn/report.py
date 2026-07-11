# Author: Ali Alaqoul <alialaqoul@gmail.com>
"""Assessment report generation — export an AI job as Markdown / HTML / DOCX.

Builds a professional penetration-test report (executive summary, finding
summary, scope, testing process, per-finding detail with evidence, prioritised
recommendations, and a command-log appendix) entirely from the session
transcript, the results store, and the deterministic analysis — nothing is
hardcoded to any environment. DOCX uses python-docx when installed;
Markdown/HTML always work with no dependencies.
"""
from __future__ import annotations

import html as _html
import re as _re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

_SEV_ORDER = ["Critical", "High", "Medium", "Low", "Info"]


def _routable_hosts(hosts: dict) -> dict:
    """Drop loopback / localhost entries — never real assessment targets."""
    out = {}
    for h, entry in (hosts or {}).items():
        hl = str(h).strip().lower()
        if hl.startswith("127.") or hl in ("::1", "localhost", "0.0.0.0"):
            continue
        out[h] = entry
    return out


def _clean_summary(text: str) -> str:
    """Turn the model's final message into clean prose for the exec summary.

    Strips the machine-readable EVIDENCE block the agent appends to its
    narrative, drops a repeated 'Executive Summary' label, and normalises
    markdown emphasis / bullet glyphs so they don't render literally.
    """
    t = (text or "").strip()
    if not t:
        return ""
    for marker in ("--- EVIDENCE", "EVIDENCE (from tool", "--- Discovered",
                   "Discovered variables:"):
        i = t.find(marker)
        if i != -1:
            t = t[:i].strip()
    t = _re.sub(r"^\**\s*executive summary\s*\**[:\-]?\s*", "", t, flags=_re.I)
    t = t.replace("**", "").replace("__", "")
    # Models often write bullets inline ("include: * a * b" / "paths: • x • y").
    # Break those onto their own lines so they render as a real list.
    t = _re.sub(r"\s*[•‣▪]\s+", "\n- ", t)                 # glyph bullets
    t = _re.sub(r"(?<=[\w\).:,])\s+\*\s+", "\n- ", t)       # ' * ' between words
    t = _re.sub(r"^\s*[•‣▪·*]\s+", "- ", t, flags=_re.M)    # line-start bullet
    t = _re.sub(r"[•‣▪·]", "-", t)                          # any stray glyph
    t = _re.sub(r"\n{3,}", "\n\n", t)
    return t.strip()


@dataclass
class Engagement:
    engagement: str = "Security assessment"
    client: str = ""
    assessor: str = ""
    authorized_by: str = ""
    date: str = field(default_factory=lambda: datetime.now().strftime("%Y-%m-%d"))
    target: str = ""
    objective: str = ""

    def rows(self):
        return [("Engagement", self.engagement), ("Client", self.client),
                ("Assessor", self.assessor), ("Authorized by", self.authorized_by),
                ("Date", self.date), ("Target", self.target),
                ("Objective", self.objective)]


_METHODOLOGY = (
    "Testing followed a standard methodology: reconnaissance and service "
    "discovery, targeted enumeration of each exposed service, analysis of the "
    "results, and identification of realistic attack paths. All actions were run "
    "against the authorized in-scope targets and are recorded in the command-log "
    "appendix. Testing was non-destructive; where a technique was validated it "
    "was taken only to a proof-of-concept."
)


def _tool_purposes():
    try:
        from .tools.registry import default_registry
        return {t.name: (t.description or "").split(".")[0]
                for t in default_registry(include_unavailable=True).all()}
    except Exception:
        return {}


# Internal signal facts used to drive logic — not shown as "variables".
_SIGNAL_FACTS = {"smb_signing", "smb_nullauth", "smb_guest", "pwned",
                 "has_users", "kerberoastable", "asreproastable"}
_SECRET_FACTS = {"password", "nthash"}


def _variables(facts: dict) -> list:
    """Meaningful engagement variables (incl. anything a custom action harvested),
    excluding internal signal facts; secrets are masked."""
    rows = []
    for k, v in (facts or {}).items():
        if k in _SIGNAL_FACTS or not v:
            continue
        val = "•" * min(10, len(str(v))) if k in _SECRET_FACTS else str(v)
        rows.append({"name": k, "value": val})
    rows.sort(key=lambda r: r["name"])
    return rows


# ATT&CK tactic order for the coverage section (kill-chain order).
_ATTACK_TACTIC_ORDER = ["reconnaissance", "resource-development", "initial-access",
    "execution", "persistence", "privilege-escalation", "stealth", "defense-impairment",
    "defense-evasion", "credential-access", "discovery", "lateral-movement", "collection",
    "command-and-control", "exfiltration", "impact"]


def build_model(meta: Engagement, transcript: list, hosts: dict,
                facts: dict, final: str, log_dir=None) -> dict:
    from .analysis import assess, attack_path, build_findings, extract_results

    hosts = _routable_hosts(hosts)
    analysis = assess(hosts, facts or {})
    findings = build_findings(hosts, facts or {}, transcript, log_dir)
    # Credentials/users are taken from actual tool output (the transcript), not
    # from the transient username/password facts (which mutate during a run).
    results = extract_results(transcript)

    # Severity counts.
    counts = {s: 0 for s in _SEV_ORDER}
    for f in findings:
        counts[f["severity"]] = counts.get(f["severity"], 0) + 1

    # Scope: one row per live host with role and OS.
    role_by_host = {h["host"]: h["role"] for h in analysis["hosts"]}
    scope = []
    for host, entry in sorted(hosts.items()):
        open_ports = [p for p in entry.get("ports", {}).values()
                      if p.get("state") == "open"]
        if not open_ports:
            continue
        os_ = entry.get("facts", {}).get("os") or ""
        scope.append({"ip": host, "hostname": entry.get("hostname") or "",
                      "role": role_by_host.get(host, "Host") or "", "os": os_})

    # Tools used (unique, in first-use order) with a one-line purpose.
    purposes = _tool_purposes()
    seen, tools_used = set(), []
    actions = [e for e in transcript if e.get("kind") == "tool_result"]
    for e in actions:
        n = e.get("name")
        if n and n not in seen:
            seen.add(n)
            tools_used.append({"tool": n, "purpose": purposes.get(n, "")})

    # Prioritised recommendations grouped by severity.
    sev_priority = {"Critical": "High", "High": "High", "Medium": "Medium",
                    "Low": "Low", "Info": "Low"}
    recs = []
    for f in findings:
        recs.append({"priority": sev_priority.get(f["severity"], "Low"),
                     "action": f["recommendation"], "finding": f["id"]})
    recs.sort(key=lambda r: {"High": 0, "Medium": 1, "Low": 2}[r["priority"]])

    # Executive summary: prefer the model's real narrative (cleaned of the
    # machine evidence block and markdown noise), else a factual auto-summary.
    cleaned = _clean_summary(final)
    exec_summary = cleaned if _real(cleaned) else _auto_summary(analysis, counts)

    # MITRE ATT&CK coverage (techniques exercised), ordered by kill-chain tactic
    # then confirmed-first — shared by all three report formats.
    from . import attack as _attack
    coverage = _attack.coverage(findings, transcript)
    _torder = {t: i for i, t in enumerate(_ATTACK_TACTIC_ORDER)}
    coverage.sort(key=lambda r: (_torder.get(r["tactic"], 99), not r["confirmed"],
                                 r["technique"]))
    from . import store as _store
    try:
        _services = _store.service_matrix()
    except Exception:
        _services = []
    attack_gaps = _attack.gaps(coverage, _services)
    path = attack_path(transcript, findings, facts or {})
    from . import bloodhound as _bh
    try:
        bh = _bh.analyze([str(log_dir) if log_dir else ".", "."], (facts or {}).get("username"))
    except Exception:
        bh = {"collected": False}

    return {
        "meta": meta, "exec_summary": exec_summary, "analysis": analysis,
        "findings": findings, "counts": counts, "scope": scope,
        "tools_used": tools_used, "methodology": _METHODOLOGY,
        "recommendations": recs, "attack": coverage,
        "attack_gaps": attack_gaps, "attack_path": path, "bloodhound": bh,
        "variables": _variables(facts or {}),
        "credentials": results["credentials"], "users": results["users"],
        "command_log": [{"tool": e.get("name", ""),
                         "command": e.get("command", "") or f"{e.get('name')} {e.get('args', {})}",
                         "ok": e.get("ok", False),
                         "output": (e.get("output") or e.get("summary") or "")}
                        for e in actions
                        if "127.0.0.1" not in (e.get("command", "") + str(e.get("args", "")))],
        "generated": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }


def _auto_summary(analysis: dict, counts: dict) -> str:
    parts = [analysis.get("summary", "").strip()]
    tot = sum(counts.values())
    if tot:
        by = ", ".join(f"{counts[s]} {s.lower()}" for s in _SEV_ORDER if counts[s])
        parts.append(f"{tot} finding(s) were identified ({by}).")
    for h in analysis.get("hosts", []):
        if h.get("attack_paths"):
            parts.append("The most likely path to compromise starts from "
                         f"{h['host']} ({h['role']}).")
            break
    return " ".join(p for p in parts if p)


def _real(text: str) -> bool:
    t = (text or "").strip()
    if not t or len(t) < 20:
        return False
    if t.startswith("{") and "action" in t[:40]:
        return False
    return not t.startswith("Assessment complete")


# ---- Markdown ---------------------------------------------------------------

def to_markdown(m: dict) -> str:
    meta: Engagement = m["meta"]
    o = ["# Penetration Test Report", "", f"## {meta.engagement}", ""]
    for k, v in meta.rows():
        if v and k != "Engagement":
            o.append(f"- **{k}:** {v}")

    o += ["", "## 1. Executive Summary", "", m["exec_summary"] or "_None._", ""]
    if m.get("attack_path"):
        o += ["### 1.1 Attack Path", "",
              "The assessment progressed to domain compromise through the following chain:", ""]
        for stp in m["attack_path"]:
            o.append(f"{stp['n']}. **{stp['title']}** — {stp['detail']}")
        o.append("")
    _bh = m.get("bloodhound") or {}
    if _bh.get("collected"):
        o += ["### 1.2 Escalation Opportunities (BloodHound)", "", _bh.get("summary", ""), ""]
        _rec = _bh.get("recommendation")
        if _rec:
            o += [f"**Recommended next escalation:** {_rec['action']} — `{_rec['first_edge']}` on "
                  f"`{_rec['path_to']}`" + (f" (tool: {_rec['tool']})" if _rec.get("tool") else ""), ""]
        if _bh.get("foothold_control"):
            o += ["| Target | Right | Abuse |", "|---|---|---|"]
            for x in _bh["foothold_control"][:12]:
                o.append(f"| {x['target']} | {x['right']} | {x['action']} |")
            o.append("")

    o += ["## 2. Finding Summary", "", "| Severity | Count |", "|---|---|"]
    for s in _SEV_ORDER:
        o.append(f"| {s} | {m['counts'][s]} |")
    o += ["", "| ID | Title | Severity | Host(s) |", "|---|---|---|---|"]
    for f in m["findings"]:
        o.append(f"| {f['id']} | {f['title']} | {f['severity']} | "
                 f"{', '.join(f['hosts'])} |")
    if not m["findings"]:
        o.append("| — | No findings identified | — | — |")

    o += ["", "## 3. Scope Overview", "",
          "| IP Address | Hostname | Role | OS |", "|---|---|---|---|"]
    for s in m["scope"]:
        o.append(f"| {s['ip']} | {s['hostname']} | {s['role']} | {s['os']} |")

    # Recovered access + variables — reflects what every action (incl. custom
    # actions/playbooks) actually harvested this engagement.
    if m.get("credentials") or m.get("users") or m.get("variables"):
        o += ["", "## Recovered Access & Discovered Variables", ""]
        if m.get("credentials"):
            o += ["**Recovered credentials**", "", "| Username | Secret | Domain | Source |",
                  "|---|---|---|---|"]
            for c in m["credentials"]:
                o.append(f"| {c['username']} | {c.get('password') or '(hash)'} | "
                         f"{c.get('domain', '')} | {c.get('note', '')} |")
            o.append("")
        if m.get("users"):
            o += ["**Enumerated users**", "", ", ".join(m["users"]), ""]
        if m.get("variables"):
            o += ["**Captured variables**", "", "| Variable | Value |", "|---|---|"]
            for v in m["variables"]:
                o.append(f"| {v['name']} | {v['value']} |")
            o.append("")

    o += ["", "## 4. Testing Process", "", "### 4.1 Methodology", "",
          m["methodology"], "", "### 4.2 Tools Used", "",
          "| Tool | Purpose |", "|---|---|"]
    for t in m["tools_used"]:
        o.append(f"| {t['tool']} | {t['purpose']} |")

    o += ["", "## 5. Findings", ""]
    for f in m["findings"]:
        _cv = f" · CVSS {f['cvss']}" if f.get("cvss") else ""
        o += [f"### {f['id']} — {f['title']} ({f['severity']}{_cv})", ""]
        if f.get("attack"):
            o += [f"**MITRE ATT&CK:** {', '.join(f['attack'])}", ""]
        o += ["**Description**", "", f["description"], ""]
        if f["evidence_cmd"]:
            o += ["**Evidence — Command**", "", "```", f["evidence_cmd"], "```", ""]
        if f["evidence_out"]:
            o += ["**Evidence — Output**", "", "```", f["evidence_out"][:1200], "```", ""]
        o += ["**Impact**", "", f["impact"], "", "**Recommendation**", "",
              f["recommendation"], ""]
    if not m["findings"]:
        o += ["_No findings were identified in this assessment._", ""]

    # ATT&CK coverage
    cov = m.get("attack") or []
    conf = sum(1 for r in cov if r["confirmed"])
    tac = len({r["tactic"] for r in cov if r["tactic"]})
    o += ["", "## 6. MITRE ATT&CK Coverage", ""]
    if cov:
        o += [f"This assessment exercised **{len(cov)} technique(s)** across **{tac} tactic(s)** "
              f"(**{conf} confirmed** by a finding, **{len(cov) - conf} attempted**). Import the "
              "JSON layer (Console → ATT&CK → *Download Navigator layer*) into the MITRE "
              "ATT&CK Navigator or VECTR for the full matrix.", "",
              "| Tactic | Technique | Name | Status |", "|---|---|---|---|"]
        for r in cov:
            o.append(f"| {(r['tactic'] or '').replace('-', ' ').title()} | {r['technique']} "
                     f"| {r['name']} | {'Confirmed' if r['confirmed'] else 'Attempted'} |")
        o.append("")
        confirmed_rows = [r for r in cov if r["confirmed"]]
        if confirmed_rows:
            o += ["### 6.1 Detection Guidance", "",
                  "What a defender should have observed for each confirmed technique:", "",
                  "| Technique | Data sources | Detection guidance |", "|---|---|---|"]
            for r in confirmed_rows:
                det = r.get("detection") or {}
                o.append(f"| {r['technique']} — {r['name']} | {det.get('data_sources', '')} "
                         f"| {det.get('guidance', '')} |")
            o.append("")
    else:
        o += ["_No ATT&CK techniques were exercised in this assessment._", ""]
    _gaps = m.get("attack_gaps") or []
    if _gaps:
        o += ["### 6.2 Coverage Gaps — Applicable but Untested", "",
              "Techniques this environment is exposed to that were not confirmed — recommended "
              "as the next round of testing (and detection engineering):", "",
              "| Tactic | Technique | Name |", "|---|---|---|"]
        for g in _gaps:
            o.append(f"| {(g['tactic'] or '').replace('-', ' ').title()} | {g['technique']} "
                     f"| {g['name']} |")
        o.append("")

    o += ["## 7. Recommendations (Prioritised)", "",
          "| Priority | Action | Finding |", "|---|---|---|"]
    for r in m["recommendations"]:
        o.append(f"| {r['priority']} | {r['action']} | {r['finding']} |")

    o += ["", "## Appendix A — Command Log", ""]
    for c in m["command_log"]:
        o += [f"**{c['tool']}** — `{c['command']}`", "", "```",
              c["output"][:1000].strip(), "```", ""]
    o += [f"_Generated by Autopwn on {m['generated']} — authorized testing only._"]
    return "\n".join(o)


# ---- HTML -------------------------------------------------------------------

_CSS = """
body{font-family:'Segoe UI',Arial,sans-serif;color:#1a1a1a;margin:40px;line-height:1.5}
h1{color:#0b5394;border-bottom:3px solid #0b5394;padding-bottom:6px}
h2{color:#0b5394;margin-top:26px;border-bottom:1px solid #ccc;padding-bottom:3px}
h3{color:#222;margin-top:18px}h4{color:#444;margin:10px 0 2px}
table{border-collapse:collapse;width:100%;margin:8px 0;font-size:12px;table-layout:fixed}
th,td{border:1px solid #ccc;padding:6px 8px;text-align:left;vertical-align:top;word-wrap:break-word}
th{background:#0b5394;color:#fff}tbody tr:nth-child(even){background:#f4f7fb}
tr{page-break-inside:avoid}thead{display:table-header-group}
.sev-Critical{color:#fff;background:#7b0000;padding:1px 5px;white-space:nowrap}.sev-High{color:#fff;background:#c00;padding:1px 5px;white-space:nowrap}
.sev-Medium{color:#000;background:#f4b400;padding:1px 5px;white-space:nowrap}.sev-Low{color:#fff;background:#3c78d8;padding:1px 5px;white-space:nowrap}
.sev-Info{color:#fff;background:#888;padding:1px 5px;white-space:nowrap}
pre{background:#f5f5f5;border:1px solid #ddd;padding:8px;font-size:11px;white-space:pre-wrap;word-wrap:break-word;overflow-wrap:break-word}
.meta td{border:none;padding:2px 8px}.foot{color:#888;font-size:11px;margin-top:30px;border-top:1px solid #ddd;padding-top:8px}
"""


def _e(s):
    return _html.escape(str(s or ""))


def _pre(text: str, width: int = 88) -> str:
    """Escape text for a <pre> block and hard-wrap over-long lines.

    Long unbreakable tokens (e.g. a comma-separated port list) can overflow the
    page when the HTML is printed. We insert real breaks at a separator near the
    width — preserving all existing whitespace (nmap output columns) rather than
    collapsing it the way textwrap would.
    """
    out = []
    for line in str(text or "").split("\n"):
        while len(line) > width:
            seg = line[:width]
            cut = max(seg.rfind(","), seg.rfind(" "), seg.rfind("/"),
                      seg.rfind(";"), seg.rfind("|"))
            if cut < width - 24:      # no good separator near the edge → hard cut
                cut = width - 1
            out.append(line[:cut + 1])
            line = line[cut + 1:]
        out.append(line)
    return _e("\n".join(out))


def _summary_html(text: str) -> str:
    """Render cleaned summary text as HTML paragraphs and bullet lists."""
    para: list[str] = []
    bullets: list[str] = []
    out: list[str] = []

    def flush_para():
        if para:
            out.append(f"<p>{_e(' '.join(para))}</p>")
            para.clear()

    def flush_bullets():
        if bullets:
            out.append("<ul>" + "".join(f"<li>{_e(b)}</li>" for b in bullets) + "</ul>")
            bullets.clear()

    for ln in (text or "").splitlines():
        s = ln.strip()
        if not s:
            flush_para(); flush_bullets(); continue
        if s.startswith("- "):
            flush_para(); bullets.append(s[2:].strip())
        else:
            flush_bullets(); para.append(s)
    flush_para(); flush_bullets()
    return "".join(out) or f"<p>{_e(text)}</p>"


def to_html(m: dict) -> str:
    meta: Engagement = m["meta"]
    p = [f"<html><head><meta charset='utf-8'><style>{_CSS}</style></head><body>"]
    p.append(f"<h1>Penetration Test Report</h1><h2 style='border:none'>{_e(meta.engagement)}</h2>")
    p.append("<table class='meta'>")
    for k, v in meta.rows():
        if v and k != "Engagement":
            p.append(f"<tr><td><b>{_e(k)}</b></td><td>{_e(v)}</td></tr>")
    p.append("</table>")

    p.append("<h2>1. Executive Summary</h2>")
    p.append(_summary_html(m['exec_summary']))
    if m.get("attack_path"):
        p.append("<h3>1.1 Attack Path</h3>")
        p.append("<p>The assessment progressed to domain compromise through the following chain:</p><ol>")
        for stp in m["attack_path"]:
            p.append(f"<li><b>{_e(stp['title'])}</b> — {_e(stp['detail'])}</li>")
        p.append("</ol>")
    _bh = m.get("bloodhound") or {}
    if _bh.get("collected"):
        p.append("<h3>1.2 Escalation Opportunities (BloodHound)</h3>")
        p.append(f"<p>{_e(_bh.get('summary', ''))}</p>")
        _rec = _bh.get("recommendation")
        if _rec:
            p.append(f"<p><b>Recommended next escalation:</b> {_e(_rec['action'])} — "
                     f"<code>{_e(_rec['first_edge'])}</code> on <code>{_e(_rec['path_to'])}</code>"
                     + (f" (tool: {_e(_rec['tool'])})" if _rec.get("tool") else "") + "</p>")
        if _bh.get("foothold_control"):
            p.append("<table><thead><tr><th width='30%'>Target</th><th width='22%'>Right</th>"
                     "<th width='48%'>Abuse</th></tr></thead><tbody>")
            for x in _bh["foothold_control"][:12]:
                p.append(f"<tr><td><code>{_e(x['target'])}</code></td><td>{_e(x['right'])}</td>"
                         f"<td>{_e(x['action'])}</td></tr>")
            p.append("</tbody></table>")

    p.append("<h2>2. Finding Summary</h2>")
    p.append("<table><thead><tr><th width='60%'>Severity</th><th width='40%'>Count</th></tr></thead><tbody>")
    for s in _SEV_ORDER:
        p.append(f"<tr><td width='60%'><span class='sev-{s}'>{s}</span></td>"
                 f"<td width='40%'>{m['counts'][s]}</td></tr>")
    p.append("</tbody></table>")
    p.append("<table><thead><tr><th width='8%'>ID</th><th width='40%'>Title</th>"
             "<th width='16%'>Severity</th><th width='36%'>Host(s)</th></tr></thead><tbody>")
    for f in m["findings"]:
        p.append(f"<tr><td width='8%'>{f['id']}</td><td width='40%'>{_e(f['title'])}</td>"
                 f"<td width='16%'><span class='sev-{f['severity']}'>{f['severity']}</span></td>"
                 f"<td width='36%'>{_e(', '.join(f['hosts']))}</td></tr>")
    if not m["findings"]:
        p.append("<tr><td>—</td><td>No findings identified</td><td>—</td><td>—</td></tr>")
    p.append("</tbody></table>")

    p.append("<h2>3. Scope Overview</h2>")
    p.append("<table><thead><tr><th width='20%'>IP Address</th><th width='22%'>Hostname</th>"
             "<th width='28%'>Role</th><th width='30%'>OS</th></tr></thead><tbody>")
    for s in m["scope"]:
        p.append(f"<tr><td width='20%'>{_e(s['ip'])}</td><td width='22%'>{_e(s['hostname'])}</td>"
                 f"<td width='28%'>{_e(s['role'])}</td><td width='30%'>{_e(s['os'])}</td></tr>")
    p.append("</tbody></table>")

    if m.get("credentials") or m.get("users") or m.get("variables"):
        p.append("<h2>Recovered Access &amp; Discovered Variables</h2>")
        if m.get("credentials"):
            p.append("<h3>Recovered credentials</h3>")
            p.append("<table><thead><tr><th>Username</th><th>Secret</th>"
                     "<th>Domain</th><th>Source</th></tr></thead><tbody>")
            for c in m["credentials"]:
                p.append(f"<tr><td>{_e(c['username'])}</td><td>{_e(c.get('password') or '(hash)')}</td>"
                         f"<td>{_e(c.get('domain', ''))}</td><td>{_e(c.get('note', ''))}</td></tr>")
            p.append("</tbody></table>")
        if m.get("users"):
            p.append("<h3>Enumerated users</h3><p>" + ", ".join(_e(u) for u in m["users"]) + "</p>")
        if m.get("variables"):
            p.append("<h3>Captured variables</h3>")
            p.append("<table><thead><tr><th width='30%'>Variable</th><th>Value</th></tr></thead><tbody>")
            for v in m["variables"]:
                p.append(f"<tr><td width='30%'>{_e(v['name'])}</td><td>{_e(v['value'])}</td></tr>")
            p.append("</tbody></table>")

    p.append("<h2>4. Testing Process</h2><h3>4.1 Methodology</h3>")
    p.append(f"<p>{_e(m['methodology'])}</p><h3>4.2 Tools Used</h3>")
    p.append("<table><thead><tr><th width='25%'>Tool</th><th width='75%'>Purpose</th></tr></thead><tbody>")
    for t in m["tools_used"]:
        p.append(f"<tr><td width='25%'>{_e(t['tool'])}</td><td width='75%'>{_e(t['purpose'])}</td></tr>")
    p.append("</tbody></table>")

    p.append("<h2>5. Findings</h2>")
    for f in m["findings"]:
        _cv = f" <span class='badge'>CVSS {_e(f['cvss'])}</span>" if f.get("cvss") else ""
        p.append(f"<h3>{f['id']} — {_e(f['title'])} "
                 f"<span class='sev-{f['severity']}'>{f['severity']}</span>{_cv}</h3>")
        if f.get("attack"):
            _att = " ".join(f"<span class='badge'>{_e(t)}</span>" for t in f["attack"])
            p.append(f"<p class='attack'><b>MITRE ATT&amp;CK:</b> {_att}</p>")
        p.append(f"<h4>Description</h4><p>{_e(f['description'])}</p>")
        if f["evidence_cmd"]:
            p.append(f"<h4>Evidence — Command</h4><pre>{_pre(f['evidence_cmd'])}</pre>")
        if f["evidence_out"]:
            p.append(f"<h4>Evidence — Output</h4><pre>{_pre(f['evidence_out'][:1200])}</pre>")
        p.append(f"<h4>Impact</h4><p>{_e(f['impact'])}</p>")
        p.append(f"<h4>Recommendation</h4><p>{_e(f['recommendation'])}</p>")
    if not m["findings"]:
        p.append("<p><i>No findings were identified in this assessment.</i></p>")

    cov = m.get("attack") or []
    p.append("<h2>6. MITRE ATT&amp;CK Coverage</h2>")
    if cov:
        conf = sum(1 for r in cov if r["confirmed"])
        tac = len({r["tactic"] for r in cov if r["tactic"]})
        p.append(f"<p>This assessment exercised <b>{len(cov)}</b> technique(s) across <b>{tac}</b> "
                 f"tactic(s) — <b>{conf}</b> confirmed by a finding, <b>{len(cov) - conf}</b> "
                 "attempted. Import the JSON layer (Console &rarr; ATT&amp;CK &rarr; "
                 "<i>Download Navigator layer</i>) into the MITRE ATT&amp;CK Navigator or VECTR "
                 "for the full matrix.</p>")
        p.append("<table><thead><tr><th width='22%'>Tactic</th><th width='14%'>Technique</th>"
                 "<th width='48%'>Name</th><th width='16%'>Status</th></tr></thead><tbody>")
        for r in cov:
            tac_lbl = _e((r["tactic"] or "").replace("-", " ").title())
            pill = ("<span style='background:#007A3D;color:#fff;padding:1px 8px;border-radius:10px;"
                    "font-size:11px'>Confirmed</span>" if r["confirmed"] else
                    "<span style='background:#D0EFDF;color:#0f172a;padding:1px 8px;border-radius:10px;"
                    "font-size:11px'>Attempted</span>")
            p.append(f"<tr><td>{tac_lbl}</td><td><code>{_e(r['technique'])}</code></td>"
                     f"<td>{_e(r['name'])}</td><td>{pill}</td></tr>")
        p.append("</tbody></table>")
        _cr = [r for r in cov if r["confirmed"]]
        if _cr:
            p.append("<h3>6.1 Detection Guidance</h3>")
            p.append("<p>What a defender should have observed for each confirmed technique:</p>")
            p.append("<table><thead><tr><th width='24%'>Technique</th><th width='30%'>Data sources</th>"
                     "<th width='46%'>Detection guidance</th></tr></thead><tbody>")
            for r in _cr:
                det = r.get("detection") or {}
                p.append(f"<tr><td><code>{_e(r['technique'])}</code> {_e(r['name'])}</td>"
                         f"<td>{_e(det.get('data_sources', ''))}</td>"
                         f"<td>{_e(det.get('guidance', ''))}</td></tr>")
            p.append("</tbody></table>")
    else:
        p.append("<p><i>No ATT&amp;CK techniques were exercised in this assessment.</i></p>")
    _gaps = m.get("attack_gaps") or []
    if _gaps:
        p.append("<h3>6.2 Coverage Gaps — Applicable but Untested</h3>")
        p.append("<p>Techniques this environment is exposed to that were not confirmed — "
                 "recommended as the next round of testing and detection engineering:</p>")
        p.append("<table><thead><tr><th width='24%'>Tactic</th><th width='16%'>Technique</th>"
                 "<th width='60%'>Name</th></tr></thead><tbody>")
        for g in _gaps:
            p.append(f"<tr><td>{_e((g['tactic'] or '').replace('-', ' ').title())}</td>"
                     f"<td><code>{_e(g['technique'])}</code></td><td>{_e(g['name'])}</td></tr>")
        p.append("</tbody></table>")

    p.append("<h2>7. Recommendations (Prioritised)</h2>")
    p.append("<table><thead><tr><th width='14%'>Priority</th><th width='72%'>Action</th>"
             "<th width='14%'>Finding</th></tr></thead><tbody>")
    for r in m["recommendations"]:
        p.append(f"<tr><td width='14%'>{r['priority']}</td><td width='72%'>{_e(r['action'])}</td>"
                 f"<td width='14%'>{r['finding']}</td></tr>")
    p.append("</tbody></table>")

    p.append("<h2>Appendix A — Command Log</h2>")
    for c in m["command_log"]:
        p.append(f"<h4>{_e(c['tool'])} — <code>{_pre(c['command'])}</code></h4>"
                 f"<pre>{_pre(c['output'][:1000].strip())}</pre>")
    p.append(f"<div class='foot'>Generated by Autopwn on {m['generated']} — "
             "for authorized security testing only.</div></body></html>")
    return "\n".join(p)


# ---- DOCX -------------------------------------------------------------------
# Palette mirrored from the HTML report so both look the same.
_DX_BLUE = (0x0B, 0x53, 0x94)          # headings + table header background
_DX_WHITE = (0xFF, 0xFF, 0xFF)
_DX_SEV = {                            # severity -> (fill hex, text rgb) — matches CSS
    "Critical": ("7B0000", _DX_WHITE), "High": ("C00000", _DX_WHITE),
    "Medium": ("F4B400", (0, 0, 0)), "Low": ("3C78D8", _DX_WHITE),
    "Info": ("888888", _DX_WHITE),
}


def _dx_shade(cell, fill_hex: str) -> None:
    """Set a table cell's background fill (like a CSS background-color)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _dx_cell(cell, text, bold=False, rgb=None, size=None, mono=False) -> None:
    from docx.shared import Pt, RGBColor
    cell.text = ""
    run = cell.paragraphs[0].add_run(_xml_safe(str(text)))
    run.font.name = "Consolas" if mono else "Arial"
    if bold:
        run.bold = True
    if rgb is not None:
        run.font.color.rgb = RGBColor(*rgb)
    if size:
        run.font.size = Pt(size)


def _dx_table(doc, headers, widths=None):
    """A Table-Grid table with a blue header row + white bold text (HTML look)."""
    from docx.shared import Inches
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    for i, h in enumerate(headers):
        _dx_cell(t.rows[0].cells[i], h, bold=True, rgb=_DX_WHITE)
        _dx_shade(t.rows[0].cells[i], "0B5394")
        if widths and i < len(widths):
            t.rows[0].cells[i].width = Inches(widths[i])
    return t


def _force_arial(doc) -> None:
    """Make Arial the document-wide font and colour headings like the HTML."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import RGBColor
    for name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3",
                 "Heading 4", "List Bullet"):
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        st.font.name = "Arial"
        rpr = st.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
        for a in ("w:ascii", "w:hAnsi", "w:cs"):
            rfonts.set(qn(a), "Arial")
    for name, rgb in (("Title", _DX_BLUE), ("Heading 1", _DX_BLUE),
                      ("Heading 2", (0x22, 0x22, 0x22)), ("Heading 3", (0x44, 0x44, 0x44))):
        try:
            doc.styles[name].font.color.rgb = RGBColor(*rgb)
        except KeyError:
            pass


def to_docx(m: dict, path: Path) -> bool:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except Exception:
        return False
    try:
        meta: Engagement = m["meta"]
        doc = Document()
        _force_arial(doc)
        doc.add_heading("Penetration Test Report", level=0)
        doc.add_heading(meta.engagement, level=1)
        t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"
        for k, v in meta.rows():
            if v and k != "Engagement":
                r = t.add_row().cells
                _dx_cell(r[0], k, bold=True); _dx_shade(r[0], "F4F7FB")
                _dx_cell(r[1], v)

        doc.add_heading("1. Executive Summary", level=1)
        for ln in (m["exec_summary"] or "").splitlines():
            s = ln.strip()
            if not s:
                continue
            if s.startswith("- "):
                doc.add_paragraph(s[2:].strip(), style="List Bullet")
            else:
                doc.add_paragraph(s)
        if m.get("attack_path"):
            doc.add_heading("1.1 Attack Path", level=2)
            doc.add_paragraph("The assessment progressed to domain compromise through the "
                              "following chain:")
            for stp in m["attack_path"]:
                doc.add_paragraph(f"{stp['title']} — {stp['detail']}", style="List Number")
        _bh = m.get("bloodhound") or {}
        if _bh.get("collected"):
            doc.add_heading("1.2 Escalation Opportunities (BloodHound)", level=2)
            doc.add_paragraph(_bh.get("summary", ""))
            _rec = _bh.get("recommendation")
            if _rec:
                doc.add_paragraph(f"Recommended next escalation: {_rec['action']} — "
                                  f"{_rec['first_edge']} on {_rec['path_to']}"
                                  + (f" (tool: {_rec['tool']})" if _rec.get("tool") else ""))
            if _bh.get("foothold_control"):
                bt = _dx_table(doc, ["Target", "Right", "Abuse"])
                for x in _bh["foothold_control"][:12]:
                    c = bt.add_row().cells
                    _dx_cell(c[0], x["target"]); _dx_cell(c[1], x["right"])
                    _dx_cell(c[2], x["action"])

        doc.add_heading("2. Finding Summary", level=1)
        st = _dx_table(doc, ["Severity", "Count"])
        for s in _SEV_ORDER:
            c = st.add_row().cells
            fill, txt = _DX_SEV[s]
            _dx_cell(c[0], s, bold=True, rgb=txt); _dx_shade(c[0], fill)
            _dx_cell(c[1], m["counts"][s])
        ft = _dx_table(doc, ["ID", "Title", "Severity", "Host(s)"])
        for f in m["findings"]:
            c = ft.add_row().cells
            _dx_cell(c[0], f["id"]); _dx_cell(c[1], f["title"])
            fill, txt = _DX_SEV.get(f["severity"], ("888888", _DX_WHITE))
            _dx_cell(c[2], f["severity"], bold=True, rgb=txt); _dx_shade(c[2], fill)
            _dx_cell(c[3], ", ".join(f["hosts"]))
        if not m["findings"]:
            c = ft.add_row().cells
            for i, v in enumerate(["—", "No findings identified", "—", "—"]):
                _dx_cell(c[i], v)

        doc.add_heading("3. Scope Overview", level=1)
        sc = _dx_table(doc, ["IP Address", "Hostname", "Role", "OS"])
        for s in m["scope"]:
            c = sc.add_row().cells
            _dx_cell(c[0], s["ip"]); _dx_cell(c[1], s["hostname"])
            _dx_cell(c[2], s["role"]); _dx_cell(c[3], s["os"])

        if m.get("credentials") or m.get("users") or m.get("variables"):
            doc.add_heading("Recovered Access & Discovered Variables", level=1)
            if m.get("credentials"):
                doc.add_heading("Recovered credentials", level=2)
                ct = _dx_table(doc, ["Username", "Secret", "Domain", "Source"])
                for cr in m["credentials"]:
                    c = ct.add_row().cells
                    _dx_cell(c[0], cr["username"]); _dx_cell(c[1], cr.get("password") or "(hash)")
                    _dx_cell(c[2], cr.get("domain", "")); _dx_cell(c[3], cr.get("note", ""))
            if m.get("users"):
                doc.add_heading("Enumerated users", level=2)
                doc.add_paragraph(", ".join(m["users"]))
            if m.get("variables"):
                doc.add_heading("Captured variables", level=2)
                vt = _dx_table(doc, ["Variable", "Value"])
                for v in m["variables"]:
                    c = vt.add_row().cells
                    _dx_cell(c[0], v["name"]); _dx_cell(c[1], v["value"])

        doc.add_heading("4. Testing Process", level=1)
        doc.add_heading("4.1 Methodology", level=2)
        doc.add_paragraph(m["methodology"])
        doc.add_heading("4.2 Tools Used", level=2)
        tt = _dx_table(doc, ["Tool", "Purpose"])
        for t2 in m["tools_used"]:
            c = tt.add_row().cells
            _dx_cell(c[0], t2["tool"]); _dx_cell(c[1], t2["purpose"])

        doc.add_heading("5. Findings", level=1)
        for f in m["findings"]:
            h = doc.add_heading("", level=2)
            h.add_run(f"{f['id']} — {f['title']}  ")
            fill, txt = _DX_SEV.get(f["severity"], ("888888", _DX_WHITE))
            badge = h.add_run(f" {f['severity']} ")
            badge.bold = True; badge.font.color.rgb = RGBColor(*txt)
            _dx_run_shade(badge, fill)
            if f.get("cvss"):
                h.add_run(f"  CVSS {f['cvss']}").italic = True
            if f.get("attack"):
                ap = doc.add_paragraph("MITRE ATT&CK: " + ", ".join(f["attack"]))
                ap.runs[0].italic = True
            doc.add_heading("Description", level=3)
            doc.add_paragraph(f["description"])
            if f["evidence_cmd"]:
                doc.add_heading("Evidence — Command", level=3)
                _mono(doc, f["evidence_cmd"])
            if f["evidence_out"]:
                doc.add_heading("Evidence — Output", level=3)
                _mono(doc, f["evidence_out"][:1200])
            doc.add_heading("Impact", level=3)
            doc.add_paragraph(f["impact"])
            doc.add_heading("Recommendation", level=3)
            doc.add_paragraph(f["recommendation"])
        if not m["findings"]:
            doc.add_paragraph("No findings were identified in this assessment.")

        doc.add_heading("6. MITRE ATT&CK Coverage", level=1)
        cov = m.get("attack") or []
        if cov:
            conf = sum(1 for r in cov if r["confirmed"])
            tac = len({r["tactic"] for r in cov if r["tactic"]})
            doc.add_paragraph(
                f"This assessment exercised {len(cov)} technique(s) across {tac} tactic(s) — "
                f"{conf} confirmed by a finding, {len(cov) - conf} attempted. Import the JSON "
                "layer (Console → ATT&CK → Download Navigator layer) into the MITRE "
                "ATT&CK Navigator or VECTR for the full matrix.")
            at = _dx_table(doc, ["Tactic", "Technique", "Name", "Status"])
            for r in cov:
                c = at.add_row().cells
                _dx_cell(c[0], (r["tactic"] or "").replace("-", " ").title())
                _dx_cell(c[1], r["technique"]); _dx_cell(c[2], r["name"])
                if r["confirmed"]:
                    _dx_cell(c[3], "Confirmed", bold=True, rgb=_DX_WHITE)
                    _dx_shade(c[3], "007A3D")
                else:
                    _dx_cell(c[3], "Attempted")
            _cr = [r for r in cov if r["confirmed"]]
            if _cr:
                doc.add_heading("6.1 Detection Guidance", level=2)
                doc.add_paragraph("What a defender should have observed for each confirmed technique:")
                dt = _dx_table(doc, ["Technique", "Data sources", "Detection guidance"])
                for r in _cr:
                    det = r.get("detection") or {}
                    c = dt.add_row().cells
                    _dx_cell(c[0], f"{r['technique']} — {r['name']}")
                    _dx_cell(c[1], det.get("data_sources", "")); _dx_cell(c[2], det.get("guidance", ""))
        else:
            doc.add_paragraph("No ATT&CK techniques were exercised in this assessment.")
        _gaps = m.get("attack_gaps") or []
        if _gaps:
            doc.add_heading("6.2 Coverage Gaps — Applicable but Untested", level=2)
            doc.add_paragraph("Techniques this environment is exposed to that were not confirmed — "
                              "recommended as the next round of testing and detection engineering:")
            gt = _dx_table(doc, ["Tactic", "Technique", "Name"])
            for g in _gaps:
                c = gt.add_row().cells
                _dx_cell(c[0], (g["tactic"] or "").replace("-", " ").title())
                _dx_cell(c[1], g["technique"]); _dx_cell(c[2], g["name"])

        doc.add_heading("7. Recommendations (Prioritised)", level=1)
        rt = _dx_table(doc, ["Priority", "Action", "Finding"])
        for r in m["recommendations"]:
            c = rt.add_row().cells
            _dx_cell(c[0], r["priority"]); _dx_cell(c[1], r["action"])
            _dx_cell(c[2], r["finding"])

        doc.add_heading("Appendix A — Command Log", level=1)
        for c in m["command_log"]:
            doc.add_heading(f"{c['tool']}", level=3)
            _mono(doc, c["command"])
            if c["output"].strip():
                _mono(doc, c["output"][:1000].strip())

        doc.add_paragraph()
        foot = doc.add_paragraph(
            f"Generated by Autopwn on {m['generated']} — for authorized "
            "security testing only.")
        foot.runs[0].font.name = "Arial"
        foot.runs[0].font.size = Pt(8)
        foot.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        doc.save(str(path))
        return True
    except Exception:
        import os as _os
        if _os.environ.get("AUTOPWN_DEBUG"):
            import traceback
            traceback.print_exc()
        return False


# Strip ANSI escapes and control characters that are invalid in XML/DOCX.
_ANSI = _re.compile(r"\x1b\[[0-9;]*[A-Za-z]")
_CTRL = _re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")


def _xml_safe(text: str) -> str:
    return _CTRL.sub("", _ANSI.sub("", str(text or "")))


def _dx_run_shade(run, fill_hex: str) -> None:
    """Give a run a background highlight (the severity 'badge' look)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    rpr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), fill_hex)
    rpr.append(shd)


def _mono(doc, text: str):
    """Evidence/command block: monospace in a light-grey bordered box, mirroring
    the HTML report's <pre> styling (so tool output stays column-aligned)."""
    from docx.shared import Pt
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    cell.text = ""
    run = cell.paragraphs[0].add_run(_xml_safe(text))
    run.font.name = "Consolas"; run.font.size = Pt(8)
    _dx_shade(cell, "F5F5F5")


# ---- export -----------------------------------------------------------------

def export(m: dict, base: Path, formats: list[str]) -> list[Path]:
    written: list[Path] = []
    if "md" in formats:
        p = base.with_suffix(".md"); p.write_text(to_markdown(m), encoding="utf-8")
        written.append(p)
    if "html" in formats:
        p = base.with_suffix(".html"); p.write_text(to_html(m), encoding="utf-8")
        written.append(p)
    if "docx" in formats:
        p = base.with_suffix(".docx")
        if to_docx(m, p):
            written.append(p)
    return written
